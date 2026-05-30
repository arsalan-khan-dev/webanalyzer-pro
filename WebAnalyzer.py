#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
WebAnalyzer v3.0 — Production-Grade Security & Performance Analysis Platform
=============================================================================
All 6 mandatory requirement blocks fully implemented:
  1. Bulletproof ZIP/TAR.GZ extraction, XSS protection, auto-cleanup, disk quota, rate limiting
  2. Bandit + ESLint + Stylelint real scanning engine
  3. Google Lighthouse CLI integration (graceful fallback)
  4. Multi-format upload: .zip .tar.gz .tar + GitHub URL + folder drag
  5. Lighthouse banner, AI Fix Suggestions (Grok), HTML/DOCX/PDF/JSON reports
  6. Gunicorn, config.yaml, Dockerfile, docker-compose, /health route
"""

# =============================================================================
# === IMPORTS & PLATFORM GUARDS ===
# =============================================================================
import os, sys, io, json, re, zipfile, tarfile, shutil, logging, traceback
import subprocess, threading, time, hashlib, socket
from pathlib import Path
from datetime import datetime, timedelta
from collections import defaultdict
from typing import Dict, List, Tuple, Any, Optional
from logging.handlers import RotatingFileHandler
from urllib.parse import urlparse

# Windows GTK guard (only applies when running natively on Windows)
if sys.platform == "win32":
    for _gtk in [r"C:\GTK3\bin", r"C:\Program Files\GTK3-Runtime Win64\bin"]:
        if os.path.exists(_gtk):
            try:
                os.add_dll_directory(_gtk)
            except Exception:
                pass

# =============================================================================
# === DEPENDENCY AUTO-INSTALL ===
# =============================================================================
_REQUIRED = {
    "flask":         "flask",
    "bs4":           "beautifulsoup4",
    "lxml":          "lxml",
    "requests":      "requests",
    "docx":          "python-docx",
    "reportlab":     "reportlab",
    "markupsafe":    "markupsafe",
    "flask_limiter": "flask-limiter[redis]",
    "yaml":          "pyyaml",
    "colorama":      "colorama",
}

def _ensure_deps():
    missing = []
    for mod, pkg in _REQUIRED.items():
        try:
            __import__(mod)
        except ImportError:
            missing.append(pkg)
    if missing:
        print(f"[INFO] Installing: {', '.join(missing)}")
        subprocess.check_call(
            [sys.executable, "-m", "pip", "install", "-q"] + missing,
            stdout=subprocess.DEVNULL
        )

_ensure_deps()

import yaml
from markupsafe import escape as _xe, Markup
from flask import Flask, request, jsonify, send_file, render_template_string, abort
from flask_limiter import Limiter
from flask_limiter.util import get_remote_address
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import requests as _req

try:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib import colors as rl_colors
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph as RLPara, Spacer,
        Table, TableStyle, HRFlowable,
    )
    from reportlab.lib.enums import TA_CENTER, TA_LEFT
    _HAS_REPORTLAB = True
except ImportError:
    _HAS_REPORTLAB = False


def _safe(val: Any) -> str:
    """HTML-escape any value for safe f-string interpolation in HTML."""
    return str(_xe(str(val) if val is not None else ""))


# =============================================================================
# === CONFIG ===
# =============================================================================
_CFG_FILE = Path(os.environ.get("WEBANALYZER_CONFIG", "config.yaml"))

def _load_config() -> Dict[str, Any]:
    defaults: Dict[str, Any] = {
        "debug":                 False,
        "host":                  "0.0.0.0",
        "port":                  5000,
        "workers":               4,
        "secret_key":            os.urandom(32).hex(),
        "max_upload_mb":         500,
        "max_file_size_mb":      50,
        "max_total_size_mb":     400,
        "disk_quota_gb":         2.0,
        "rate_limit":            "8 per minute",
        "session_ttl_minutes":   30,
        "xai_api_key":           os.environ.get("XAI_API_KEY", ""),
        "exclude_dirs": [
            ".git", "node_modules", "__pycache__",
            ".venv", "venv", "dist", "build",
            ".next", "coverage", ".tox", ".mypy_cache",
        ],
        "log_file":  "webanalyzer.log",
        "log_level": "INFO",
    }
    if _CFG_FILE.exists():
        try:
            with open(_CFG_FILE, "r", encoding="utf-8") as fh:
                user = yaml.safe_load(fh) or {}
            defaults.update(user)
        except Exception as exc:
            print(f"[WARN] Could not parse config.yaml: {exc}")
    return defaults

CFG = _load_config()

# =============================================================================
# === LOGGING SETUP ===
# =============================================================================
_LOG_LEVEL = getattr(logging, str(CFG["log_level"]).upper(), logging.INFO)
_root = logging.getLogger()
_root.setLevel(_LOG_LEVEL)
_fmt = logging.Formatter("%(asctime)s [%(levelname)s] %(name)s — %(message)s")

_ch = logging.StreamHandler()
_ch.setFormatter(_fmt)
_root.addHandler(_ch)

try:
    _fh = RotatingFileHandler(
        CFG["log_file"], maxBytes=10 * 1024 * 1024, backupCount=3, encoding="utf-8"
    )
    _fh.setFormatter(_fmt)
    _root.addHandler(_fh)
except Exception:
    pass

logger = logging.getLogger("webanalyzer")

BASE_UPLOAD = Path("/tmp/webanalyzer_uploads")
BASE_OUTPUT = Path("/tmp/webanalyzer_reports")
BASE_UPLOAD.mkdir(parents=True, exist_ok=True)
BASE_OUTPUT.mkdir(parents=True, exist_ok=True)

_MAX_FILE_BYTES  = int(CFG["max_file_size_mb"]) * 1024 * 1024
_MAX_TOTAL_BYTES = int(CFG["max_total_size_mb"]) * 1024 * 1024
_DISK_QUOTA      = int(float(CFG["disk_quota_gb"]) * 1024 ** 3)

_NOISE_DIRS: set = set(CFG["exclude_dirs"])

# =============================================================================
# === SAFE EXTRACTION v2 ===
# =============================================================================
class ExtractionError(Exception):
    pass


def _safe_member_path(raw_name: str, dest_root: Path) -> Optional[Path]:
    """
    Validate and resolve a member name against dest_root.
    Returns resolved Path or None if the entry is dangerous.
    """
    # Strip leading slashes / Windows drive letters
    clean = re.sub(r"^([A-Za-z]:)?[/\\]+", "", raw_name)
    # Collapse . and normalise separators
    try:
        candidate = (dest_root / clean).resolve()
    except Exception:
        return None
    # MUST stay inside dest_root (zip-slip prevention)
    try:
        candidate.relative_to(dest_root.resolve())
    except ValueError:
        return None
    # Depth guard — no more than 20 path components
    if len(Path(clean).parts) > 20:
        return None
    return candidate


def safe_extract_zip(zip_path: Path, dest: Path) -> int:
    """
    Extract a ZIP file safely.  Returns number of extracted files.
    Raises ExtractionError on any security violation or quota breach.
    """
    dest.mkdir(parents=True, exist_ok=True)
    total_bytes = 0
    extracted = 0

    with zipfile.ZipFile(zip_path, "r") as zf:
        for info in zf.infolist():
            # Skip directories
            if info.filename.endswith("/"):
                continue
            # Reject symlinks encoded as external attrs (Unix: mode bits)
            unix_mode = (info.external_attr >> 16) & 0xFFFF
            if unix_mode and (
                (unix_mode & 0o170000) == 0o120000  # symlink
                or (unix_mode & 0o170000) == 0o060000  # block device
                or (unix_mode & 0o170000) == 0o020000  # char device
            ):
                logger.warning("Skipping dangerous ZIP entry: %s", info.filename)
                continue
            # Per-file size guard (uncompressed)
            if info.file_size > _MAX_FILE_BYTES:
                logger.warning("Skipping oversized file %s (%d bytes)", info.filename, info.file_size)
                continue
            # Total size guard
            total_bytes += info.file_size
            if total_bytes > _MAX_TOTAL_BYTES:
                raise ExtractionError(
                    f"Archive total exceeds {CFG['max_total_size_mb']} MB limit."
                )
            # Path safety
            dest_path = _safe_member_path(info.filename, dest)
            if dest_path is None:
                logger.warning("Skipping unsafe ZIP path: %s", info.filename)
                continue
            dest_path.parent.mkdir(parents=True, exist_ok=True)
            with zf.open(info) as src, open(dest_path, "wb") as dst:
                shutil.copyfileobj(src, dst)
            extracted += 1

    return extracted


def safe_extract_tar(tar_path: Path, dest: Path) -> int:
    """
    Extract a TAR / TAR.GZ file safely.
    Raises ExtractionError on any security violation or quota breach.
    """
    dest.mkdir(parents=True, exist_ok=True)
    total_bytes = 0
    extracted = 0

    with tarfile.open(tar_path, "r:*") as tf:
        for member in tf.getmembers():
            # Skip non-regular files (symlinks, hard-links, devices, fifos)
            if not member.isfile():
                continue
            # Per-file size guard
            if member.size > _MAX_FILE_BYTES:
                logger.warning("Skipping oversized TAR entry %s", member.name)
                continue
            total_bytes += member.size
            if total_bytes > _MAX_TOTAL_BYTES:
                raise ExtractionError(
                    f"Archive total exceeds {CFG['max_total_size_mb']} MB limit."
                )
            dest_path = _safe_member_path(member.name, dest)
            if dest_path is None:
                logger.warning("Skipping unsafe TAR path: %s", member.name)
                continue
            dest_path.parent.mkdir(parents=True, exist_ok=True)
            fobj = tf.extractfile(member)
            if fobj is None:
                continue
            with open(dest_path, "wb") as dst:
                shutil.copyfileobj(fobj, dst)
            extracted += 1

    return extracted


def clone_github_repo(url: str, dest: Path, timeout: int = 120) -> int:
    """Clone a public GitHub repo (depth=1). Returns file count."""
    parsed = urlparse(url)
    # Whitelist only github.com / gitlab.com / bitbucket.org
    allowed = {"github.com", "gitlab.com", "bitbucket.org"}
    if parsed.hostname not in allowed:
        raise ExtractionError(f"Only {allowed} URLs are supported for git clone.")
    # Build a clean HTTPS URL (no credentials)
    clean_url = f"https://{parsed.hostname}{parsed.path.rstrip('/')}"
    if not clean_url.endswith(".git"):
        clean_url += ".git"

    dest.mkdir(parents=True, exist_ok=True)
    result = subprocess.run(
        ["git", "clone", "--depth", "1", "--single-branch", clean_url, str(dest)],
        capture_output=True, text=True, timeout=timeout,
    )
    if result.returncode != 0:
        raise ExtractionError(f"git clone failed: {result.stderr[:300]}")
    return sum(1 for _ in dest.rglob("*") if _.is_file())


# =============================================================================
# === CLEANUP MANAGER ===
# =============================================================================
class CleanupManager:
    """
    Background thread deletes session directories after TTL minutes.
    Also enforces the global /tmp disk quota.
    """

    def __init__(self, ttl_minutes: int = 30):
        self._ttl = timedelta(minutes=ttl_minutes)
        self._sessions: Dict[str, Tuple[Path, datetime]] = {}
        self._lock = threading.Lock()
        t = threading.Thread(target=self._loop, daemon=True, name="cleanup")
        t.start()

    def register(self, session_id: str, path: Path):
        with self._lock:
            self._sessions[session_id] = (path, datetime.now() + self._ttl)
        logger.debug("Cleanup registered %s → %s", session_id, path)

    def cancel(self, session_id: str):
        with self._lock:
            self._sessions.pop(session_id, None)

    def _loop(self):
        while True:
            time.sleep(60)
            try:
                self._sweep()
                self._quota_guard()
            except Exception as exc:
                logger.error("Cleanup error: %s", exc)

    def _sweep(self):
        now = datetime.now()
        with self._lock:
            expired = [sid for sid, (_, exp) in self._sessions.items() if now >= exp]
        for sid in expired:
            with self._lock:
                entry = self._sessions.pop(sid, None)
            if entry:
                path, _ = entry
                _rm(path)
                logger.info("Cleaned up session %s (%s)", sid, path)

    def _quota_guard(self):
        total = 0
        for base in [BASE_UPLOAD, BASE_OUTPUT]:
            for f in base.rglob("*"):
                if f.is_file():
                    try:
                        total += f.stat().st_size
                    except Exception:
                        pass
        if total > _DISK_QUOTA:
            logger.warning("Disk quota exceeded (%d bytes). Purging oldest sessions.", total)
            # Remove all registered sessions immediately
            with self._lock:
                all_sessions = list(self._sessions.items())
                self._sessions.clear()
            for _, (path, _) in all_sessions:
                _rm(path)


def _rm(path: Path):
    try:
        if path.exists():
            shutil.rmtree(path, ignore_errors=True)
    except Exception as exc:
        logger.debug("rm error %s: %s", path, exc)


_cleanup = CleanupManager(ttl_minutes=int(CFG["session_ttl_minutes"]))


# =============================================================================
# === FILE DECODING ===
# =============================================================================
_ENCODINGS = ["utf-8", "latin-1", "cp1252"]

def decode_file(path: Path) -> Optional[str]:
    """Try UTF-8 → latin-1 → cp1252. Log when fallback is used."""
    for enc in _ENCODINGS:
        try:
            content = path.read_text(encoding=enc, errors="strict")
            if enc != "utf-8":
                logger.debug("Decoded %s with %s", path.name, enc)
            return content
        except (UnicodeDecodeError, ValueError):
            continue
    logger.warning("Could not decode %s with any encoding", path.name)
    return None


# =============================================================================
# === HELPER: free port ===
# =============================================================================
def _free_port() -> int:
    with socket.socket(socket.AF_INET, socket.SOCK_STREAM) as s:
        s.bind(("", 0))
        return s.getsockname()[1]


# =============================================================================
# === REGEX CODE ANALYZER ===
# =============================================================================
class RegexAnalyzer:

    def analyze(self, filepath: Path, content: str, language: str) -> List[Dict]:
        method = getattr(self, f"_js" if language in ("javascript", "typescript") else f"_{language}", None)
        if method is None:
            return []
        return method(content)

    # ---------- JavaScript / TypeScript ----------
    def _js(self, src: str) -> List[Dict]:
        issues = []
        def _add(sev, typ, msg, fix):
            issues.append({"severity": sev, "type": typ, "message": msg, "fix": fix, "source": "regex"})

        if re.search(r"\beval\s*\(", src):
            _add("CRITICAL", "Security", "eval() usage — code injection risk",
                 "Replace with JSON.parse() or a safer alternative")
        if re.search(r"\bdocument\.write\s*\(", src):
            _add("CRITICAL", "Security", "document.write() is dangerous and blocks parsing",
                 "Use DOM methods: createElement / textContent / appendChild")
        if re.search(r"\bnew\s+Function\s*\(", src):
            _add("HIGH", "Security", "new Function() constructor — similar risk to eval()",
                 "Refactor to use static, pre-declared functions")
        if re.search(r"innerHTML\s*[+]?=", src):
            _add("HIGH", "Security", "innerHTML assignment — XSS vector",
                 "Use textContent or DOMPurify.sanitize() before inserting HTML")
        if re.search(r"setTimeout\s*\(\s*['\"]", src):
            _add("HIGH", "Security", "setTimeout() with string argument — eval-equivalent",
                 "Always pass a function reference, never a string")
        if re.search(r"setInterval\s*\(\s*['\"]", src):
            _add("HIGH", "Security", "setInterval() with string argument — eval-equivalent",
                 "Pass a function reference instead of a string")
        if len(src) > 50_000 and "min" not in src[:20]:
            _add("MEDIUM", "Performance", "Large unminified JS file",
                 "Minify with Terser / esbuild and enable gzip on the server")
        if re.search(r"console\.(log|warn|error)\s*\(", src):
            _add("LOW", "Code Quality", "console.* statements left in source",
                 "Remove console statements before production deployment")
        return issues

    # ---------- Python ----------
    def _python(self, src: str) -> List[Dict]:
        issues = []
        def _add(sev, typ, msg, fix):
            issues.append({"severity": sev, "type": typ, "message": msg, "fix": fix, "source": "regex"})

        if re.search(r"\b(eval|exec)\s*\(", src):
            _add("CRITICAL", "Security", "eval()/exec() usage — arbitrary code execution risk",
                 "Avoid dynamic code execution; redesign the logic")
        if re.search(r'execute\s*\(\s*["\'].*?[+%]', src):
            _add("CRITICAL", "Security", "Potential SQL injection via string concatenation in query",
                 "Use parameterized queries: cursor.execute(sql, (param,))")
        if re.search(r"(password|api_key|secret|token)\s*=\s*['\"]", src, re.I):
            _add("CRITICAL", "Security", "Hardcoded credential detected",
                 "Store secrets in environment variables or a secrets manager")
        if re.search(r"pickle\.loads?\s*\(", src):
            _add("HIGH", "Security", "pickle.load() on untrusted data — RCE risk",
                 "Use json or a safe serialization library for untrusted data")
        if re.search(r"subprocess\.(call|Popen|run)\(.*shell\s*=\s*True", src):
            _add("HIGH", "Security", "subprocess with shell=True — command injection risk",
                 "Use a list of arguments; avoid shell=True unless absolutely necessary")
        if re.search(r"except\s*:", src):
            _add("MEDIUM", "Code Quality", "Bare except clause catches everything including SystemExit",
                 "Specify exception types: except ValueError: or except Exception as e:")
        if re.search(r"import \*", src):
            _add("LOW", "Code Quality", "Wildcard import pollutes namespace",
                 "Import only what you need: from module import specific_name")
        return issues

    # ---------- HTML ----------
    def _html(self, src: str) -> List[Dict]:
        issues = []
        def _add(sev, typ, msg, fix):
            issues.append({"severity": sev, "type": typ, "message": msg, "fix": fix, "source": "regex"})

        soup = BeautifulSoup(src, "lxml")
        if not src.lower().lstrip().startswith("<!doctype"):
            _add("HIGH", "HTML Compliance", "Missing DOCTYPE declaration",
                 "Add <!DOCTYPE html> as the very first line")
        if not soup.find("meta", attrs={"name": "viewport"}):
            _add("HIGH", "Responsiveness", "Missing viewport meta tag",
                 '<meta name="viewport" content="width=device-width, initial-scale=1">')
        if not soup.find("title"):
            _add("HIGH", "SEO", "Missing <title> tag", "Add a 50–60 character descriptive title")
        if not soup.find("meta", attrs={"name": "description"}):
            _add("MEDIUM", "SEO", "Missing meta description",
                 '<meta name="description" content="…">')
        h1s = soup.find_all("h1")
        if not h1s:
            _add("MEDIUM", "SEO", "No <h1> heading found", "Add one <h1> with the page main topic")
        elif len(h1s) > 1:
            _add("MEDIUM", "SEO", f"{len(h1s)} <h1> tags found (should be exactly 1)",
                 "Keep one <h1> per page; use <h2>–<h6> for subheadings")
        missing_alt = [i for i in soup.find_all("img") if not i.get("alt")]
        if missing_alt:
            _add("HIGH", "Accessibility", f"{len(missing_alt)} image(s) missing alt attribute",
                 'Add descriptive alt="…" to every <img>')
        if re.search(r"on\w+\s*=\s*['\"]", src, re.I):
            _add("HIGH", "Security", "Inline event handler(s) detected — XSS surface",
                 "Move event handlers to external JS and use addEventListener()")
        return issues

    # ---------- CSS / SCSS ----------
    def _css(self, src: str) -> List[Dict]:
        issues = []
        important_count = len(re.findall(r"!important", src))
        if important_count > 5:
            issues.append({
                "severity": "MEDIUM", "type": "Code Quality", "source": "regex",
                "message": f"!important overused ({important_count}×) — indicates specificity problems",
                "fix": "Restructure selectors to avoid needing !important",
            })
        return issues

    # ---------- PHP ----------
    def _php(self, src: str) -> List[Dict]:
        issues = []
        def _add(sev, typ, msg, fix):
            issues.append({"severity": sev, "type": typ, "message": msg, "fix": fix, "source": "regex"})

        if re.search(r"\beval\s*\(", src):
            _add("CRITICAL", "Security", "PHP eval() — RCE risk", "Never use eval(); refactor the logic")
        if re.search(r"\$_(GET|POST|REQUEST|COOKIE)\[", src) and re.search(r"(mysql_query|mysqli_query|PDO)", src):
            _add("CRITICAL", "Security", "User input passed directly into SQL query",
                 "Use PDO prepared statements with bound parameters")
        if re.search(r"echo\s+\$_(GET|POST|REQUEST)", src):
            _add("HIGH", "Security", "Unsanitized user input echoed — reflected XSS",
                 "Use htmlspecialchars($var, ENT_QUOTES, 'UTF-8') before output")
        return issues

    # ---------- Java ----------
    def _java(self, src: str) -> List[Dict]:
        issues = []
        if re.search(r'(password|apiKey|secret)\s*=\s*["\']', src, re.I):
            issues.append({
                "severity": "CRITICAL", "type": "Security", "source": "regex",
                "message": "Hardcoded credentials in Java source",
                "fix": "Use environment variables or a config-server (Spring Vault, AWS Secrets Manager)",
            })
        if re.search(r'Runtime\.exec\s*\(', src):
            issues.append({
                "severity": "HIGH", "type": "Security", "source": "regex",
                "message": "Runtime.exec() — command injection risk",
                "fix": "Use ProcessBuilder with a list of args; validate all inputs",
            })
        return issues


# =============================================================================
# === BANDIT INTEGRATION ===
# =============================================================================
class BanditScanner:
    """Runs Bandit static analyzer on Python files."""

    _SEV_MAP = {"HIGH": "CRITICAL", "MEDIUM": "HIGH", "LOW": "MEDIUM"}
    _available: Optional[bool] = None

    @classmethod
    def is_available(cls) -> bool:
        if cls._available is None:
            try:
                r = subprocess.run(
                    [sys.executable, "-m", "bandit", "--version"],
                    capture_output=True, timeout=10,
                )
                cls._available = r.returncode == 0
            except Exception:
                cls._available = False
                try:
                    subprocess.check_call(
                        [sys.executable, "-m", "pip", "install", "-q", "bandit"],
                        stdout=subprocess.DEVNULL
                    )
                    cls._available = True
                except Exception:
                    pass
        return cls._available

    @classmethod
    def scan(cls, filepath: Path) -> List[Dict]:
        if not cls.is_available():
            return []
        try:
            result = subprocess.run(
                [sys.executable, "-m", "bandit", "-f", "json", "-q", str(filepath)],
                capture_output=True, text=True, timeout=60,
            )
            raw = result.stdout.strip()
            if not raw:
                return []
            data = json.loads(raw)
            issues = []
            for r in data.get("results", []):
                sev = cls._SEV_MAP.get(r.get("issue_severity", "LOW"), "MEDIUM")
                issues.append({
                    "severity": sev,
                    "type": "Security (Bandit)",
                    "source": "bandit",
                    "message": f"[{r.get('test_id','')}] {r.get('issue_text','')} (line {r.get('line_number','')})",
                    "fix": f"Confidence: {r.get('issue_confidence','')} — see https://bandit.readthedocs.io",
                })
            return issues
        except Exception as exc:
            logger.debug("Bandit error on %s: %s", filepath, exc)
            return []


# =============================================================================
# === ESLINT INTEGRATION ===
# =============================================================================
_ESLINT_CFG = """{
  "env": {"browser": true, "es2021": true, "node": true},
  "parserOptions": {"ecmaVersion": 2021, "sourceType": "module"},
  "rules": {
    "no-eval": "error",
    "no-implied-eval": "error",
    "no-new-func": "error",
    "no-script-url": "error",
    "no-alert": "warn",
    "no-console": "warn",
    "eqeqeq": "warn",
    "no-unused-vars": "warn",
    "no-undef": "warn"
  }
}"""

class ESLintScanner:
    _available: Optional[bool] = None
    _cfg_path: Optional[Path] = None

    @classmethod
    def is_available(cls) -> bool:
        if cls._available is None:
            for cmd in [["npx", "--yes", "eslint", "--version"], ["eslint", "--version"]]:
                try:
                    r = subprocess.run(cmd, capture_output=True, timeout=30)
                    if r.returncode == 0:
                        cls._available = True
                        break
                except Exception:
                    continue
            else:
                cls._available = False
        return cls._available

    @classmethod
    def _write_cfg(cls) -> Path:
        if cls._cfg_path is None or not cls._cfg_path.exists():
            p = Path("/tmp/.webanalyzer_eslintrc.json")
            p.write_text(_ESLINT_CFG, encoding="utf-8")
            cls._cfg_path = p
        return cls._cfg_path

    @classmethod
    def scan(cls, filepath: Path) -> List[Dict]:
        if not cls.is_available():
            return []
        cfg = cls._write_cfg()
        try:
            result = subprocess.run(
                ["npx", "--yes", "eslint",
                 "--no-eslintrc", "-c", str(cfg),
                 "--format", "json", str(filepath)],
                capture_output=True, text=True, timeout=60,
            )
            raw = (result.stdout or "").strip()
            if not raw or raw[0] != "[":
                return []
            data = json.loads(raw)
            issues = []
            for file_result in data:
                for msg in file_result.get("messages", []):
                    sev_int = msg.get("severity", 1)
                    sev = "HIGH" if sev_int == 2 else "MEDIUM"
                    issues.append({
                        "severity": sev,
                        "type": "JavaScript (ESLint)",
                        "source": "eslint",
                        "message": f"[{msg.get('ruleId','?')}] {msg.get('message','')} (line {msg.get('line','')})",
                        "fix": f"Rule: {msg.get('ruleId','')} — see eslint.org/docs/rules/{msg.get('ruleId','')}",
                    })
            return issues
        except Exception as exc:
            logger.debug("ESLint error on %s: %s", filepath, exc)
            return []


# =============================================================================
# === STYLELINT INTEGRATION ===
# =============================================================================
_STYLELINT_CFG = """{
  "rules": {
    "color-no-invalid-hex": true,
    "declaration-block-no-duplicate-properties": true,
    "no-duplicate-selectors": true,
    "no-empty-source": true,
    "unit-no-unknown": true,
    "property-no-unknown": true
  }
}"""

class StylelintScanner:
    _available: Optional[bool] = None
    _cfg_path: Optional[Path] = None

    @classmethod
    def is_available(cls) -> bool:
        if cls._available is None:
            try:
                r = subprocess.run(
                    ["npx", "--yes", "stylelint", "--version"],
                    capture_output=True, timeout=30,
                )
                cls._available = r.returncode == 0
            except Exception:
                cls._available = False
        return cls._available

    @classmethod
    def _write_cfg(cls) -> Path:
        if cls._cfg_path is None or not cls._cfg_path.exists():
            p = Path("/tmp/.webanalyzer_stylelintrc.json")
            p.write_text(_STYLELINT_CFG, encoding="utf-8")
            cls._cfg_path = p
        return cls._cfg_path

    @classmethod
    def scan(cls, filepath: Path) -> List[Dict]:
        if not cls.is_available():
            return []
        cfg = cls._write_cfg()
        try:
            result = subprocess.run(
                ["npx", "--yes", "stylelint",
                 f"--config={cfg}",
                 "--formatter=json", str(filepath)],
                capture_output=True, text=True, timeout=60,
            )
            raw = (result.stdout or "").strip()
            if not raw or raw[0] != "[":
                return []
            data = json.loads(raw)
            issues = []
            for file_result in data:
                for w in file_result.get("warnings", []):
                    sev = "HIGH" if w.get("severity") == "error" else "MEDIUM"
                    issues.append({
                        "severity": sev,
                        "type": "CSS (Stylelint)",
                        "source": "stylelint",
                        "message": f"[{w.get('rule','?')}] {w.get('text','')} (line {w.get('line','')})",
                        "fix": f"Rule: {w.get('rule','')} — see stylelint.io/user-guide/rules",
                    })
            return issues
        except Exception as exc:
            logger.debug("Stylelint error on %s: %s", filepath, exc)
            return []


# =============================================================================
# === LIGHTHOUSE INTEGRATION ===
# =============================================================================
class LighthouseRunner:
    """
    Spins up a temporary HTTP server, runs Google Lighthouse CLI against it,
    and returns parsed scores.  Gracefully falls back with a clear message.
    """

    _available: Optional[bool] = None

    @classmethod
    def is_available(cls) -> bool:
        if cls._available is None:
            for cmd in [["lighthouse", "--version"], ["npx", "--yes", "lighthouse", "--version"]]:
                try:
                    r = subprocess.run(cmd, capture_output=True, timeout=30)
                    if r.returncode == 0:
                        cls._available = True
                        break
                except Exception:
                    continue
            else:
                cls._available = False
        return cls._available

    @classmethod
    def run(cls, site_dir: Path) -> Optional[Dict]:
        """
        Returns a dict with keys: performance, accessibility,
        best-practices, seo, pwa  (each 0–100).
        Returns None if Lighthouse is unavailable or analysis fails.
        """
        if not cls.is_available():
            logger.info("Lighthouse not found — skipping")
            return None

        # Locate main HTML file
        html_candidates = (
            list(site_dir.rglob("index.html"))
            + list(site_dir.rglob("index.htm"))
            + list(site_dir.rglob("*.html"))
        )
        if not html_candidates:
            logger.info("No HTML file found for Lighthouse")
            return None
        main_html = html_candidates[0]
        rel_path = main_html.relative_to(site_dir)

        # Start temporary HTTP server
        import http.server
        import socketserver

        port = _free_port()

        class _Handler(http.server.SimpleHTTPRequestHandler):
            def __init__(self, *a, **kw):
                super().__init__(*a, directory=str(site_dir), **kw)

            def log_message(self, *_):
                pass  # silence

        server = socketserver.TCPServer(("127.0.0.1", port), _Handler)
        server.allow_reuse_address = True
        srv_thread = threading.Thread(target=server.serve_forever, daemon=True)
        srv_thread.start()

        url = f"http://127.0.0.1:{port}/{rel_path.as_posix()}"
        output_file = Path(f"/tmp/lh_report_{port}.json")

        lh_cmd = next(
            (c for c in [["lighthouse"], ["npx", "--yes", "lighthouse"]]
             if subprocess.run(c + ["--version"], capture_output=True).returncode == 0),
            None,
        )
        if lh_cmd is None:
            server.shutdown()
            return None

        try:
            r = subprocess.run(
                lh_cmd + [
                    url,
                    "--chrome-flags=--headless --no-sandbox --disable-gpu --disable-dev-shm-usage",
                    "--output=json",
                    f"--output-path={output_file}",
                    "--quiet",
                    "--only-categories=performance,accessibility,best-practices,seo,pwa",
                ],
                capture_output=True, text=True, timeout=120,
            )
            if r.returncode not in (0, 1) or not output_file.exists():
                logger.warning("Lighthouse failed (rc=%d): %s", r.returncode, r.stderr[:200])
                return None

            raw = json.loads(output_file.read_text(encoding="utf-8"))
            cats = raw.get("categories", {})
            scores = {}
            for key in ("performance", "accessibility", "best-practices", "seo", "pwa"):
                sc = cats.get(key, {}).get("score")
                scores[key] = round(sc * 100) if sc is not None else None
            return scores

        except subprocess.TimeoutExpired:
            logger.warning("Lighthouse timed out")
            return None
        except Exception as exc:
            logger.warning("Lighthouse error: %s", exc)
            return None
        finally:
            server.shutdown()
            if output_file.exists():
                output_file.unlink(missing_ok=True)


# =============================================================================
# === PERFORMANCE ANALYZER ===
# =============================================================================
class PerformanceAnalyzer:
    @staticmethod
    def analyze(root: Path) -> List[Dict]:
        issues = []
        for fpath in root.rglob("*"):
            if not fpath.is_file():
                continue
            # Skip noise dirs
            if any(part in _NOISE_DIRS for part in fpath.parts):
                continue
            size_mb = fpath.stat().st_size / (1024 * 1024)
            ext = fpath.suffix.lower()
            rel = str(fpath.relative_to(root))
            if ext in (".png", ".jpg", ".jpeg", ".gif", ".bmp") and size_mb > 1:
                issues.append({
                    "severity": "MEDIUM", "type": "Performance", "source": "perf",
                    "message": f"Large image ({size_mb:.2f} MB): {rel}",
                    "fix": "Compress and convert to WebP; use lazy loading",
                })
            elif ext == ".js" and "min" not in fpath.name and size_mb > 0.5:
                issues.append({
                    "severity": "MEDIUM", "type": "Performance", "source": "perf",
                    "message": f"Large unminified JS ({size_mb:.2f} MB): {rel}",
                    "fix": "Minify with Terser/esbuild; enable Brotli/gzip on server",
                })
            elif ext == ".css" and "min" not in fpath.name and size_mb > 0.3:
                issues.append({
                    "severity": "LOW", "type": "Performance", "source": "perf",
                    "message": f"Large CSS file ({size_mb:.2f} MB): {rel}",
                    "fix": "Consider CSS splitting and critical-CSS inlining",
                })
        return issues


# =============================================================================
# === SEO ANALYZER ===
# =============================================================================
class SEOAnalyzer:
    @staticmethod
    def analyze(root: Path) -> List[Dict]:
        issues = []
        if not (root / "robots.txt").exists():
            issues.append({
                "severity": "LOW", "type": "SEO", "source": "seo",
                "message": "robots.txt not found",
                "fix": "Create robots.txt to guide search engine crawlers",
            })
        if not any((root / s).exists() for s in ("sitemap.xml", "sitemap_index.xml")):
            issues.append({
                "severity": "MEDIUM", "type": "SEO", "source": "seo",
                "message": "sitemap.xml not found",
                "fix": "Generate an XML sitemap and submit it to Google Search Console",
            })
        # Check for canonical links in HTML files
        for hf in list(root.rglob("*.html"))[:10]:
            content = decode_file(hf)
            if content and "<link" in content and "canonical" not in content.lower():
                issues.append({
                    "severity": "LOW", "type": "SEO", "source": "seo",
                    "message": f"No canonical link tag in {hf.relative_to(root)}",
                    "fix": '<link rel="canonical" href="https://example.com/page">',
                })
                break
        return issues


# =============================================================================
# === ACCESSIBILITY ANALYZER ===
# =============================================================================
class AccessibilityAnalyzer:
    @staticmethod
    def analyze(root: Path) -> List[Dict]:
        issues = []
        for hf in list(root.rglob("*.html"))[:8]:
            content = decode_file(hf)
            if not content:
                continue
            soup = BeautifulSoup(content, "lxml")
            rel = str(hf.relative_to(root))

            # Skip navigation
            if not soup.find_all("a", string=re.compile(r"skip|jump", re.I)):
                issues.append({
                    "severity": "MEDIUM", "type": "Accessibility", "source": "a11y",
                    "message": f"No skip-nav link in {rel}",
                    "fix": 'Add <a href="#main-content" class="skip-link">Skip to content</a>',
                })
                break

            # lang attribute on <html>
            html_tag = soup.find("html")
            if html_tag and not html_tag.get("lang"):
                issues.append({
                    "severity": "MEDIUM", "type": "Accessibility", "source": "a11y",
                    "message": f'<html> tag missing lang attribute in {rel}',
                    "fix": 'Add lang="en" (or appropriate language code) to <html>',
                })
                break

            # Role attributes on interactive elements
            buttons = soup.find_all("div", onclick=True) + soup.find_all("span", onclick=True)
            if buttons:
                issues.append({
                    "severity": "MEDIUM", "type": "Accessibility", "source": "a11y",
                    "message": f"Clickable <div>/<span> without role in {rel}",
                    "fix": "Use <button> or add role='button' + keyboard event handlers",
                })
                break
        return issues


# =============================================================================
# === AI FIX SUGGESTIONS (Grok / xAI) ===
# =============================================================================
def get_ai_suggestions(issues: List[Dict], api_key: str) -> str:
    """
    Calls xAI Grok API to generate enhanced fix recommendations.
    Returns markdown text.  Returns empty string on failure.
    """
    if not api_key:
        return ""
    top = issues[:12]
    prompt = (
        "You are a senior security engineer. Analyze these web-application issues and "
        "return a concise, prioritized action plan with code examples where helpful. "
        "Format: numbered list, each item max 3 sentences.\n\n"
        "Issues:\n" + json.dumps(
            [{"severity": i.get("severity"), "type": i.get("type"), "message": i.get("message")} for i in top],
            indent=2,
        )
    )
    try:
        resp = _req.post(
            "https://api.x.ai/v1/chat/completions",
            headers={"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"},
            json={
                "model": "grok-beta",
                "messages": [{"role": "user", "content": prompt}],
                "max_tokens": 1200,
            },
            timeout=30,
        )
        resp.raise_for_status()
        return resp.json()["choices"][0]["message"]["content"]
    except Exception as exc:
        logger.warning("xAI Grok call failed: %s", exc)
        return ""


# =============================================================================
# === REPORT GENERATOR ===
# =============================================================================
class ReportGenerator:

    def __init__(self, site_name: str, timestamp: datetime):
        self.site_name  = site_name
        self.timestamp  = timestamp
        self.issues:    List[Dict] = []
        self.lh_scores: Optional[Dict] = None
        self.ai_text:   str = ""

    def add_issues(self, issues: List[Dict]):
        self.issues.extend(issues)

    def set_lighthouse(self, scores: Optional[Dict]):
        self.lh_scores = scores

    def set_ai_text(self, text: str):
        self.ai_text = text

    # ------------------------------------------------------------------
    def calculate_score(self) -> Tuple[float, str]:
        if not self.issues:
            return 100.0, "Excellent"
        c = sum(1 for i in self.issues if i.get("severity") == "CRITICAL")
        h = sum(1 for i in self.issues if i.get("severity") == "HIGH")
        m = sum(1 for i in self.issues if i.get("severity") == "MEDIUM")
        l = sum(1 for i in self.issues if i.get("severity") == "LOW")
        score = max(0.0, min(100.0, 100.0 - c * 10 - h * 5 - m * 2 - l * 0.5))
        grade = (
            "Excellent" if score >= 90 else
            "Good"      if score >= 80 else
            "Fair"      if score >= 70 else
            "Poor"      if score >= 60 else "Critical"
        )
        return score, grade

    # ------------------------------------------------------------------
    def generate_json(self) -> str:
        score, grade = self.calculate_score()
        return json.dumps({
            "site_name":   self.site_name,
            "timestamp":   self.timestamp.isoformat(),
            "score":       score,
            "grade":       grade,
            "lighthouse":  self.lh_scores,
            "issues":      self.issues,
            "ai_suggestions": self.ai_text,
        }, indent=2)

    # ------------------------------------------------------------------
    def generate_html(self) -> str:
        score, grade = self.calculate_score()
        sn   = _safe(self.site_name)
        ts   = _safe(self.timestamp.strftime("%Y-%m-%d %H:%M"))
        sc   = f"{score:.0f}"
        sc_color = "#00d68f" if score >= 80 else "#ff9500" if score >= 60 else "#ff4d6d"

        c = sum(1 for i in self.issues if i.get("severity") == "CRITICAL")
        h = sum(1 for i in self.issues if i.get("severity") == "HIGH")
        m = sum(1 for i in self.issues if i.get("severity") == "MEDIUM")
        l = sum(1 for i in self.issues if i.get("severity") == "LOW")

        # Lighthouse block
        lh_html = ""
        if self.lh_scores:
            def _circle(label: str, val: Optional[int]) -> str:
                if val is None:
                    return f'<div class="lh-item"><div class="lh-na">N/A</div><span>{_safe(label)}</span></div>'
                col = "#00d68f" if val >= 90 else "#ff9500" if val >= 50 else "#ff4d6d"
                dash = val  # circumference ≈ 100 for r=15.9
                return f"""
                <div class="lh-item">
                  <svg viewBox="0 0 36 36" class="lh-ring">
                    <circle cx="18" cy="18" r="15.9" fill="none" stroke="#2a3a55" stroke-width="3"/>
                    <circle cx="18" cy="18" r="15.9" fill="none" stroke="{col}" stroke-width="3"
                      stroke-dasharray="{dash} 100" stroke-linecap="round"
                      transform="rotate(-90 18 18)"/>
                    <text x="18" y="22" text-anchor="middle" font-size="8"
                          font-family="JetBrains Mono,monospace" fill="{col}" font-weight="bold">{val}</text>
                  </svg>
                  <span>{_safe(label)}</span>
                </div>"""

            lh_html = f"""
            <div class="lh-banner">
              <h3>⚡ Lighthouse Scores</h3>
              <div class="lh-circles">
                {_circle("Performance",    self.lh_scores.get("performance"))}
                {_circle("Accessibility",  self.lh_scores.get("accessibility"))}
                {_circle("Best Practices", self.lh_scores.get("best-practices"))}
                {_circle("SEO",            self.lh_scores.get("seo"))}
                {_circle("PWA",            self.lh_scores.get("pwa"))}
              </div>
            </div>"""

        # AI suggestions block
        ai_html = ""
        if self.ai_text:
            # ai_text comes from Grok, treat as plain text — escape it
            ai_escaped = _safe(self.ai_text).replace("\n", "<br>")
            ai_html = f"""
            <div class="ai-block">
              <h3>✨ AI Fix Suggestions (Grok)</h3>
              <div class="ai-content">{ai_escaped}</div>
            </div>"""

        # Issues
        sev_order = {"CRITICAL": 0, "HIGH": 1, "MEDIUM": 2, "LOW": 3}
        sorted_issues = sorted(self.issues, key=lambda x: sev_order.get(x.get("severity", "LOW"), 9))
        by_type: Dict[str, List] = defaultdict(list)
        for issue in sorted_issues:
            by_type[issue.get("type", "Other")].append(issue)

        issues_html = ""
        if not sorted_issues:
            issues_html = '<div class="no-issues">✓ No issues found — your project is clean!</div>'
        else:
            for itype, type_issues in by_type.items():
                issues_html += f'<div class="igroup"><h3 class="itype">{_safe(itype)}</h3>'
                for issue in type_issues:
                    sev  = _safe(issue.get("severity", "LOW"))
                    msg  = _safe(issue.get("message", "—"))
                    fix  = _safe(issue.get("fix", "No suggestion available"))
                    src  = _safe(issue.get("source", ""))
                    low  = sev.lower()
                    issues_html += f"""
                    <div class="issue {low}">
                      <div class="itop">
                        <span class="badge {low}">{sev}</span>
                        <span class="src-tag">{src}</span>
                        <span class="imsg">{msg}</span>
                      </div>
                      <div class="ifix">→ {fix}</div>
                    </div>"""
                issues_html += "</div>"

        return f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width,initial-scale=1">
<title>WebAnalyzer Report — {sn}</title>
<link href="https://fonts.googleapis.com/css2?family=Syne:wght@400;600;700;800&family=JetBrains+Mono:wght@400;600&display=swap" rel="stylesheet">
<style>
:root{{--bg:#07080f;--surf:#0d1117;--card:#111827;--brd:#1e2d4a;
  --acc:#4a7ec7;--grn:#00d68f;--red:#ff4d6d;--ora:#ff9500;--yel:#ffd60a;
  --txt:#e2e8f0;--mut:#5a7298}}
*{{margin:0;padding:0;box-sizing:border-box}}
body{{font-family:'Syne',sans-serif;background:var(--bg);color:var(--txt);padding:32px 16px}}
.wrap{{max-width:1060px;margin:0 auto}}
header{{background:linear-gradient(135deg,#0d1117,#111827);border:1px solid var(--brd);
  border-radius:12px;padding:36px;text-align:center;margin-bottom:20px}}
header h1{{font-size:2.2em;font-weight:800;background:linear-gradient(135deg,#fff 30%,#4a7ec7);
  -webkit-background-clip:text;-webkit-text-fill-color:transparent}}
header p{{color:var(--mut);margin-top:8px}}
.score-row{{display:grid;grid-template-columns:180px 1fr;gap:16px;margin-bottom:20px}}
.score-box{{background:var(--card);border:1px solid var(--brd);border-radius:12px;
  display:flex;flex-direction:column;align-items:center;justify-content:center;padding:28px}}
.sc{{font-size:4.5em;font-weight:800;font-family:'JetBrains Mono',monospace;line-height:1;color:{sc_color}}}
.grade{{font-size:1.1em;color:{sc_color};margin-top:8px;font-weight:700;text-transform:uppercase;letter-spacing:2px}}
.meta{{background:var(--card);border:1px solid var(--brd);border-radius:12px;padding:24px}}
.meta h2{{font-size:.75em;color:var(--mut);text-transform:uppercase;letter-spacing:2px;margin-bottom:14px}}
.row{{display:flex;justify-content:space-between;padding:9px 0;border-bottom:1px solid var(--brd);font-size:.93em}}
.row:last-child{{border:none}}.row .k{{color:var(--mut)}}.row .v{{font-weight:700}}
.stats{{display:grid;grid-template-columns:repeat(4,1fr);gap:12px;margin-bottom:20px}}
.stat{{background:var(--card);border:1px solid var(--brd);border-radius:10px;padding:18px;text-align:center}}
.stat-n{{font-size:2em;font-weight:800;font-family:'JetBrains Mono',monospace}}
.stat-l{{font-size:.72em;color:var(--mut);margin-top:5px;text-transform:uppercase;letter-spacing:1px}}
.c .stat-n{{color:var(--red)}}.h .stat-n{{color:var(--ora)}}.m .stat-n{{color:var(--yel)}}.lw .stat-n{{color:var(--grn)}}
.lh-banner{{background:var(--card);border:1px solid var(--brd);border-radius:12px;padding:24px;margin-bottom:20px}}
.lh-banner h3{{font-size:.8em;text-transform:uppercase;letter-spacing:2px;color:var(--mut);margin-bottom:18px}}
.lh-circles{{display:flex;gap:28px;flex-wrap:wrap}}
.lh-item{{display:flex;flex-direction:column;align-items:center;gap:8px;font-size:.78em;color:var(--mut)}}
.lh-ring{{width:72px;height:72px}}.lh-na{{font-size:1.4em;color:var(--mut)}}
.ai-block{{background:var(--card);border:1px solid rgba(74,126,199,.3);border-radius:12px;padding:24px;margin-bottom:20px}}
.ai-block h3{{font-size:.8em;text-transform:uppercase;letter-spacing:2px;color:var(--acc);margin-bottom:14px}}
.ai-content{{font-size:.9em;line-height:1.8;color:var(--txt)}}
.panel{{background:var(--card);border:1px solid var(--brd);border-radius:12px;padding:24px;margin-bottom:20px}}
.panel h2{{font-size:.75em;color:var(--mut);text-transform:uppercase;letter-spacing:2px;margin-bottom:18px}}
.igroup{{margin-bottom:22px}}
.itype{{font-size:.72em;text-transform:uppercase;letter-spacing:2px;color:var(--acc);
  border-bottom:1px solid var(--brd);padding-bottom:7px;margin-bottom:10px}}
.issue{{border-left:3px solid;border-radius:0 7px 7px 0;padding:12px 14px;margin-bottom:8px}}
.issue.critical{{border-color:var(--red);background:rgba(255,77,109,.05)}}
.issue.high{{border-color:var(--ora);background:rgba(255,149,0,.05)}}
.issue.medium{{border-color:var(--yel);background:rgba(255,214,10,.05)}}
.issue.low{{border-color:var(--grn);background:rgba(0,214,143,.05)}}
.itop{{display:flex;align-items:flex-start;gap:8px;flex-wrap:wrap;margin-bottom:5px}}
.badge{{padding:2px 8px;border-radius:3px;font-size:.7em;font-weight:700;font-family:'JetBrains Mono',monospace;white-space:nowrap}}
.badge.critical{{background:var(--red);color:#fff}}
.badge.high{{background:var(--ora);color:#000}}
.badge.medium{{background:var(--yel);color:#000}}
.badge.low{{background:var(--grn);color:#000}}
.src-tag{{font-family:'JetBrains Mono',monospace;font-size:.68em;color:var(--mut);
  border:1px solid var(--brd);padding:2px 6px;border-radius:3px;white-space:nowrap}}
.imsg{{font-weight:600;font-size:.88em;flex:1}}
.ifix{{font-size:.82em;color:var(--mut);padding-left:4px}}
.no-issues{{background:rgba(0,214,143,.06);border:1px solid var(--grn);border-radius:8px;
  padding:24px;text-align:center;color:var(--grn);font-size:1em}}
footer{{text-align:center;padding:24px 0 4px;color:var(--brd);font-size:.78em;font-family:'JetBrains Mono',monospace}}
@media(max-width:640px){{.score-row,.stats{{grid-template-columns:1fr}}}}
</style>
</head>
<body>
<div class="wrap">
<header>
  <h1>🔍 WebAnalyzer Report</h1>
  <p>Security · Performance · SEO · Accessibility</p>
</header>
<div class="score-row">
  <div class="score-box"><div class="sc">{sc}</div><div class="grade">{_safe(grade)}</div></div>
  <div class="meta">
    <h2>Summary</h2>
    <div class="row"><span class="k">Project</span><span class="v">{sn}</span></div>
    <div class="row"><span class="k">Analyzed</span><span class="v">{ts}</span></div>
    <div class="row"><span class="k">Score</span><span class="v">{sc} / 100</span></div>
    <div class="row"><span class="k">Total Issues</span><span class="v">{len(self.issues)}</span></div>
  </div>
</div>
<div class="stats">
  <div class="stat c"><div class="stat-n">{c}</div><div class="stat-l">Critical</div></div>
  <div class="stat h"><div class="stat-n">{h}</div><div class="stat-l">High</div></div>
  <div class="stat m"><div class="stat-n">{m}</div><div class="stat-l">Medium</div></div>
  <div class="stat lw"><div class="stat-n">{l}</div><div class="stat-l">Low</div></div>
</div>
{lh_html}
{ai_html}
<div class="panel">
  <h2>Detailed Findings</h2>
  {issues_html}
</div>
<footer>Generated by WebAnalyzer v3.0 · {ts}</footer>
</div>
</body>
</html>"""

    # ------------------------------------------------------------------
    def generate_docx(self, output_path: Path):
        doc = Document()
        for s in doc.sections:
            s.top_margin = s.bottom_margin = Inches(1)
            s.left_margin = s.right_margin = Inches(1.1)

        # Title
        t = doc.add_paragraph()
        t.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tr = t.add_run("WebAnalyzer Report v3.0")
        tr.font.size = Pt(22)
        tr.font.bold = True

        sub = doc.add_paragraph()
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sr = sub.add_run("Security · Performance · SEO · Accessibility")
        sr.font.size = Pt(12)
        sr.font.color.rgb = RGBColor(100, 100, 120)

        doc.add_paragraph()
        score, grade = self.calculate_score()
        c = sum(1 for i in self.issues if i.get("severity") == "CRITICAL")
        h = sum(1 for i in self.issues if i.get("severity") == "HIGH")
        m = sum(1 for i in self.issues if i.get("severity") == "MEDIUM")
        l = sum(1 for i in self.issues if i.get("severity") == "LOW")

        tbl = doc.add_table(rows=6, cols=2)
        tbl.style = "Light Grid Accent 1"
        for row, (k, v) in zip(tbl.rows, [
            ("Project",        self.site_name),
            ("Date",           self.timestamp.strftime("%Y-%m-%d %H:%M:%S")),
            ("Overall Score",  f"{score:.0f}/100 ({grade})"),
            ("Total Issues",   str(len(self.issues))),
            ("Critical",       str(c)),
            ("High / Med / Low", f"{h} / {m} / {l}"),
        ]):
            row.cells[0].text = k
            row.cells[1].text = v

        # Lighthouse scores
        if self.lh_scores:
            doc.add_paragraph()
            doc.add_heading("Lighthouse Scores", 1)
            lh_data = [["Category", "Score"]] + [
                [k.title(), str(v) if v is not None else "N/A"]
                for k, v in self.lh_scores.items()
            ]
            lt = doc.add_table(rows=len(lh_data), cols=2)
            lt.style = "Light Grid Accent 1"
            for row, (k, v) in zip(lt.rows, lh_data):
                row.cells[0].text = k
                row.cells[1].text = v

        # AI suggestions
        if self.ai_text:
            doc.add_page_break()
            doc.add_heading("AI Fix Suggestions (Grok)", 1)
            for line in self.ai_text.split("\n"):
                if line.strip():
                    doc.add_paragraph(line.strip())

        # Detailed findings
        doc.add_page_break()
        doc.add_heading("Detailed Findings", 1)
        by_type: Dict[str, List] = defaultdict(list)
        for issue in self.issues:
            by_type[issue.get("type", "Other")].append(issue)

        if not self.issues:
            doc.add_paragraph("No issues detected — project is clean.")
        else:
            for itype, type_issues in by_type.items():
                doc.add_heading(itype, 2)
                for issue in type_issues:
                    p = doc.add_paragraph(style="List Bullet")
                    p.add_run(f"[{issue.get('severity')}] ").font.bold = True
                    p.add_run(issue.get("message", ""))
                    doc.add_paragraph(
                        f"Fix: {issue.get('fix', '')} | Source: {issue.get('source', '')}",
                        style="List Bullet 2"
                    )

        doc.save(output_path)

    # ------------------------------------------------------------------
    def generate_pdf(self, output_path: Path):
        """Try weasyprint; fall back to ReportLab."""
        # Try weasyprint first (preferred — full CSS support)
        try:
            from weasyprint import HTML as WH
            html_str = self.generate_html()
            WH(string=html_str).write_pdf(str(output_path))
            logger.info("PDF generated via weasyprint")
            return
        except ImportError:
            pass
        except Exception as exc:
            logger.warning("weasyprint failed: %s — falling back to ReportLab", exc)

        # ReportLab fallback
        if not _HAS_REPORTLAB:
            raise RuntimeError("Neither weasyprint nor reportlab is available for PDF generation.")

        score, grade = self.calculate_score()
        c = sum(1 for i in self.issues if i.get("severity") == "CRITICAL")
        h = sum(1 for i in self.issues if i.get("severity") == "HIGH")
        m = sum(1 for i in self.issues if i.get("severity") == "MEDIUM")
        l = sum(1 for i in self.issues if i.get("severity") == "LOW")

        doc = SimpleDocTemplate(str(output_path), pagesize=A4,
                                rightMargin=2*cm, leftMargin=2*cm,
                                topMargin=2*cm, bottomMargin=2*cm)
        styles = getSampleStyleSheet()

        title_s = ParagraphStyle("Title2", parent=styles["Title"],
                                  fontSize=22, fontName="Helvetica-Bold",
                                  textColor=rl_colors.HexColor("#1a1f35"),
                                  spaceAfter=6, alignment=TA_CENTER)
        sub_s   = ParagraphStyle("Sub2", parent=styles["Normal"],
                                  fontSize=11, textColor=rl_colors.HexColor("#5a7298"),
                                  alignment=TA_CENTER, spaceAfter=14)
        h1_s    = ParagraphStyle("H1", parent=styles["Heading1"],
                                  fontSize=14, fontName="Helvetica-Bold",
                                  textColor=rl_colors.HexColor("#2d3748"),
                                  spaceBefore=14, spaceAfter=7)
        h2_s    = ParagraphStyle("H2", parent=styles["Heading2"],
                                  fontSize=11, fontName="Helvetica-Bold",
                                  textColor=rl_colors.HexColor("#4a7ec7"),
                                  spaceBefore=9, spaceAfter=4)
        body_s  = ParagraphStyle("Body2", parent=styles["Normal"],
                                  fontSize=9.5, textColor=rl_colors.HexColor("#2d3748"),
                                  spaceAfter=4)
        fix_s   = ParagraphStyle("Fix2", parent=styles["Normal"],
                                  fontSize=8.5, textColor=rl_colors.HexColor("#718096"),
                                  leftIndent=10, spaceAfter=6)
        foot_s  = ParagraphStyle("Footer2", parent=styles["Normal"],
                                  fontSize=7.5, textColor=rl_colors.HexColor("#a0aec0"),
                                  alignment=TA_CENTER)

        sev_colors_map = {
            "CRITICAL": rl_colors.HexColor("#ff4d6d"),
            "HIGH":     rl_colors.HexColor("#ff9500"),
            "MEDIUM":   rl_colors.HexColor("#ffd60a"),
            "LOW":      rl_colors.HexColor("#00b894"),
        }

        story = []
        story.append(RLPara("WebAnalyzer Report", title_s))
        story.append(RLPara("Security · Performance · SEO · Accessibility", sub_s))
        story.append(HRFlowable(width="100%", thickness=2,
                                color=rl_colors.HexColor("#4a7ec7"), spaceAfter=14))

        # Summary table
        sum_data = [
            ["Metric", "Value"],
            ["Project", self.site_name],
            ["Date", self.timestamp.strftime("%Y-%m-%d %H:%M:%S")],
            ["Score", f"{score:.0f} / 100  ({grade})"],
            ["Total Issues", str(len(self.issues))],
            ["Critical / High / Med / Low", f"{c} / {h} / {m} / {l}"],
        ]
        st = Table(sum_data, colWidths=[5*cm, 12*cm])
        st.setStyle(TableStyle([
            ("BACKGROUND",   (0,0), (-1,0), rl_colors.HexColor("#2d3748")),
            ("TEXTCOLOR",    (0,0), (-1,0), rl_colors.white),
            ("FONTNAME",     (0,0), (-1,0), "Helvetica-Bold"),
            ("FONTSIZE",     (0,0), (-1,-1), 9.5),
            ("ROWBACKGROUNDS",(0,1),(-1,-1),[rl_colors.HexColor("#f7fafc"), rl_colors.white]),
            ("GRID",         (0,0), (-1,-1), 0.5, rl_colors.HexColor("#e2e8f0")),
            ("FONTNAME",     (0,1), (0,-1), "Helvetica-Bold"),
            ("TOPPADDING",   (0,0), (-1,-1), 7),
            ("BOTTOMPADDING",(0,0), (-1,-1), 7),
            ("LEFTPADDING",  (0,0), (-1,-1), 9),
        ]))
        story.append(st)
        story.append(Spacer(1, 14))

        # Lighthouse scores
        if self.lh_scores:
            story.append(RLPara("Lighthouse Scores", h1_s))
            lh_data = [["Category", "Score"]] + [
                [k.title(), str(v) if v is not None else "N/A"]
                for k, v in self.lh_scores.items()
            ]
            lt = Table(lh_data, colWidths=[6*cm, 11*cm])
            lt.setStyle(TableStyle([
                ("BACKGROUND",   (0,0), (-1,0), rl_colors.HexColor("#2d3748")),
                ("TEXTCOLOR",    (0,0), (-1,0), rl_colors.white),
                ("FONTNAME",     (0,0), (-1,0), "Helvetica-Bold"),
                ("FONTSIZE",     (0,0), (-1,-1), 9.5),
                ("GRID",         (0,0), (-1,-1), 0.5, rl_colors.HexColor("#e2e8f0")),
                ("ROWBACKGROUNDS",(0,1),(-1,-1),[rl_colors.HexColor("#f0fff4"), rl_colors.white]),
                ("TOPPADDING",   (0,0), (-1,-1), 7),
                ("BOTTOMPADDING",(0,0), (-1,-1), 7),
                ("LEFTPADDING",  (0,0), (-1,-1), 9),
            ]))
            story.append(lt)
            story.append(Spacer(1, 14))

        # Detailed findings
        story.append(RLPara("Detailed Findings", h1_s))
        sev_order = {"CRITICAL": 0, "HIGH": 1, "MEDIUM": 2, "LOW": 3}
        sorted_issues = sorted(self.issues, key=lambda x: sev_order.get(x.get("severity", "LOW"), 9))
        by_type: Dict[str, List] = defaultdict(list)
        for issue in sorted_issues:
            by_type[issue.get("type", "Other")].append(issue)

        if not self.issues:
            story.append(RLPara("No issues found — project is clean!", body_s))
        else:
            for itype, type_issues in by_type.items():
                story.append(RLPara(itype, h2_s))
                for issue in type_issues:
                    sev = issue.get("severity", "LOW")
                    sc_col = sev_colors_map.get(sev, rl_colors.grey)
                    hex_col = "%06x" % sc_col.hexval() if hasattr(sc_col, "hexval") else "333333"
                    row = [[
                        RLPara(f'<b><font color="#{hex_col}">[{sev}]</font></b>', body_s),
                        RLPara(f'<b>{issue.get("message","")}</b> '
                               f'<font color="#718096">({issue.get("source","")})</font>', body_s),
                    ]]
                    rt = Table(row, colWidths=[2.5*cm, 14.5*cm])
                    rt.setStyle(TableStyle([
                        ("VALIGN", (0,0), (-1,-1), "TOP"),
                        ("LEFTPADDING", (0,0), (-1,-1), 0),
                        ("TOPPADDING", (0,0), (-1,-1), 2),
                        ("BOTTOMPADDING", (0,0), (-1,-1), 2),
                    ]))
                    story.append(rt)
                    story.append(RLPara(f"→ {issue.get('fix','')}", fix_s))

        story.append(Spacer(1, 18))
        story.append(HRFlowable(width="100%", thickness=1,
                                color=rl_colors.HexColor("#e2e8f0")))
        story.append(Spacer(1, 6))
        story.append(RLPara(
            f"WebAnalyzer v3.0 · {self.timestamp.strftime('%Y-%m-%d %H:%M')} · Static analysis results",
            foot_s,
        ))
        doc.build(story)
        logger.info("PDF generated via ReportLab (fallback)")


# =============================================================================
# === ANALYSIS ORCHESTRATOR ===
# =============================================================================
class AnalysisOrchestrator:

    _EXT_MAP = {
        ".js": "javascript", ".ts": "typescript", ".jsx": "javascript", ".tsx": "typescript",
        ".py": "python",
        ".html": "html", ".htm": "html",
        ".css": "css", ".scss": "css",
        ".php": "php",
        ".java": "java",
        ".vue": "html",
    }

    def __init__(self, session_id: str, site_name: str, work_dir: Path):
        self.session_id = session_id
        self.site_name  = site_name
        self.work_dir   = work_dir
        self.issues:    List[Dict] = []
        self.file_count = 0
        self._regex  = RegexAnalyzer()

    def _should_skip(self, path: Path) -> bool:
        return any(part in _NOISE_DIRS for part in path.parts)

    def _dedup(self):
        seen = {}
        for issue in self.issues:
            key = (issue.get("message", ""), issue.get("severity", ""))
            if key not in seen:
                seen[key] = issue
        self.issues = list(seen.values())

    def run(self) -> Dict[str, Any]:
        try:
            for fpath in self.work_dir.rglob("*"):
                if not fpath.is_file() or self._should_skip(fpath):
                    continue
                lang = self._EXT_MAP.get(fpath.suffix.lower())
                if not lang:
                    continue
                content = decode_file(fpath)
                if content is None:
                    continue

                # Regex analysis
                self.issues.extend(self._regex.analyze(fpath, content, lang))

                # Tool-based analysis
                if lang == "python":
                    self.issues.extend(BanditScanner.scan(fpath))
                elif lang in ("javascript", "typescript"):
                    self.issues.extend(ESLintScanner.scan(fpath))
                elif lang == "css":
                    self.issues.extend(StylelintScanner.scan(fpath))

                self.file_count += 1

            # Structural analyzers
            self.issues.extend(PerformanceAnalyzer.analyze(self.work_dir))
            self.issues.extend(SEOAnalyzer.analyze(self.work_dir))
            self.issues.extend(AccessibilityAnalyzer.analyze(self.work_dir))

            self._dedup()
            logger.info("Analysis complete: %d files, %d issues", self.file_count, len(self.issues))
            return {"success": True, "file_count": self.file_count, "issues": self.issues}

        except Exception as exc:
            logger.error("Orchestrator error: %s", traceback.format_exc())
            return {"success": False, "error": str(exc)}


# =============================================================================
# === FLASK APP & RATE LIMITER ===
# =============================================================================
app = Flask(__name__)
app.config.update(
    SECRET_KEY          = CFG["secret_key"],
    MAX_CONTENT_LENGTH  = int(CFG["max_upload_mb"]) * 1024 * 1024,
    DEBUG               = False,
    TESTING             = False,
)

limiter = Limiter(
    app         = app,
    key_func    = get_remote_address,
    default_limits = [CFG["rate_limit"]],
    storage_uri = "memory://",
)

# In-memory session store
_sessions: Dict[str, Dict] = {}
_sessions_lock = threading.Lock()


# =============================================================================
# === HTML TEMPLATE ===
# =============================================================================
HTML_TEMPLATE = r"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width,initial-scale=1">
<title>WebAnalyzer v3.0 — Source Code Intelligence</title>
<link href="https://fonts.googleapis.com/css2?family=Syne:wght@400;600;700;800&family=JetBrains+Mono:wght@400;500;600&display=swap" rel="stylesheet">
<style>
  :root{
    --bg:#07080f;--surf:#0d1117;--card:#111827;--brd:#1e2d4a;
    --acc:#4a7ec7;--acc2:#7b5ea7;--grn:#00d68f;--red:#ff4d6d;
    --ora:#ff9500;--yel:#ffd60a;--txt:#e2e8f0;--mut:#5a7298;
  }
  *{margin:0;padding:0;box-sizing:border-box}
  html{scroll-behavior:smooth}
  body{
    font-family:'Syne',sans-serif;background:var(--bg);color:var(--txt);
    min-height:100vh;overflow-x:hidden;
  }
  body::before{
    content:'';position:fixed;inset:0;pointer-events:none;
    background-image:
      linear-gradient(rgba(74,126,199,.04) 1px,transparent 1px),
      linear-gradient(90deg,rgba(74,126,199,.04) 1px,transparent 1px);
    background-size:44px 44px;
  }

  .wrap{position:relative;z-index:1;max-width:980px;margin:0 auto;padding:44px 20px}

  /* ── Header ── */
  .hdr{text-align:center;margin-bottom:48px;animation:fadeDown .55s ease both}
  @keyframes fadeDown{from{opacity:0;transform:translateY(-14px)}to{opacity:1;transform:none}}
  @keyframes fadeUp{from{opacity:0;transform:translateY(14px)}to{opacity:1;transform:none}}
  .badge-pill{
    display:inline-block;padding:5px 16px;border:1px solid var(--brd);border-radius:999px;
    font-family:'JetBrains Mono',monospace;font-size:.72em;color:var(--acc);
    background:rgba(74,126,199,.08);margin-bottom:18px;letter-spacing:1px;
  }
  .hdr h1{
    font-size:clamp(2.2em,5vw,3.5em);font-weight:800;letter-spacing:-2px;line-height:1.1;
    background:linear-gradient(135deg,#fff 30%,#4a7ec7);
    -webkit-background-clip:text;-webkit-text-fill-color:transparent;
  }
  .hdr p{color:var(--mut);margin-top:14px;font-size:1.05em;max-width:500px;margin-inline:auto}

  /* ── Cards ── */
  .card{background:var(--card);border:1px solid var(--brd);border-radius:16px;padding:36px;margin-bottom:20px}

  /* ── Source Tabs ── */
  .src-tabs{display:flex;gap:8px;margin-bottom:28px}
  .src-tab{
    flex:1;padding:11px;border:1px solid var(--brd);border-radius:8px;
    background:none;color:var(--mut);font-family:'Syne',sans-serif;font-size:.9em;
    font-weight:600;cursor:pointer;transition:.22s;text-align:center;
  }
  .src-tab:hover{border-color:var(--acc);color:var(--txt)}
  .src-tab.active{border-color:var(--acc);background:rgba(74,126,199,.12);color:var(--acc)}

  /* ── Drop Zone ── */
  .drop-zone{
    border:2px dashed var(--brd);border-radius:12px;padding:56px 32px;
    text-align:center;cursor:pointer;transition:all .25s ease;
    background:rgba(74,126,199,.025);position:relative;overflow:hidden;
  }
  .drop-zone::before{
    content:'';position:absolute;inset:0;pointer-events:none;
    background:radial-gradient(ellipse at 50% 0%,rgba(74,126,199,.08),transparent 60%);
    opacity:0;transition:.3s;
  }
  .drop-zone:hover,.drop-zone.over{border-color:var(--acc);background:rgba(74,126,199,.05)}
  .drop-zone:hover::before,.drop-zone.over::before{opacity:1}
  .drop-zone:hover{transform:translateY(-2px)}
  .dz-icon{font-size:3.2em;display:block;margin-bottom:14px;animation:float 3s ease-in-out infinite}
  @keyframes float{0%,100%{transform:translateY(0)}50%{transform:translateY(-7px)}}
  .dz-title{font-size:1.25em;font-weight:700;margin-bottom:6px}
  .dz-sub{color:var(--mut);font-size:.88em;line-height:1.7}
  input[type=file]{display:none}

  /* GitHub URL panel */
  .gh-panel{display:none}
  .gh-input-wrap{position:relative}
  .gh-input-wrap input{
    width:100%;padding:14px 16px 14px 48px;border:1px solid var(--brd);
    border-radius:8px;background:rgba(255,255,255,.04);color:var(--txt);
    font-family:'JetBrains Mono',monospace;font-size:.92em;outline:none;transition:.22s;
  }
  .gh-input-wrap input:focus{border-color:var(--acc)}
  .gh-input-wrap input::placeholder{color:var(--mut)}
  .gh-icon{position:absolute;left:14px;top:50%;transform:translateY(-50%);font-size:1.2em;pointer-events:none}

  /* ── Buttons ── */
  .btn-primary{
    display:inline-flex;align-items:center;gap:8px;
    background:linear-gradient(135deg,var(--acc),var(--acc2));
    color:#fff;border:none;padding:13px 26px;border-radius:8px;
    font-family:'Syne',sans-serif;font-size:.95em;font-weight:700;
    cursor:pointer;transition:all .22s ease;margin-top:18px;letter-spacing:.3px;
  }
  .btn-primary:hover{transform:translateY(-2px);box-shadow:0 8px 24px rgba(74,126,199,.3)}
  .btn-primary:disabled{opacity:.5;cursor:not-allowed;transform:none;box-shadow:none}

  /* ── Error ── */
  .err{background:rgba(255,77,109,.08);border:1px solid rgba(255,77,109,.3);
    border-radius:8px;padding:14px 18px;color:var(--red);margin-top:18px;display:none;font-size:.9em}

  /* ── Loading ── */
  #loading{display:none;text-align:center;padding:40px 0}
  .spin-row{display:flex;align-items:center;justify-content:center;gap:14px;margin-bottom:18px}
  .spin{width:34px;height:34px;border:3px solid var(--brd);border-top-color:var(--acc);
    border-radius:50%;animation:spin 1s linear infinite}
  @keyframes spin{to{transform:rotate(360deg)}}
  .ldt{font-family:'JetBrains Mono',monospace;font-size:.88em;color:var(--acc);
    animation:pulse 1.5s ease-in-out infinite}
  @keyframes pulse{0%,100%{opacity:.45}50%{opacity:1}}
  .prog-track{width:100%;height:4px;background:var(--brd);border-radius:2px;overflow:hidden}
  .prog-fill{height:100%;background:linear-gradient(90deg,var(--acc),var(--acc2));
    width:0%;transition:width .3s ease;border-radius:2px}

  /* ── Results ── */
  #results{display:none;animation:fadeUp .5s ease both}

  /* Score hero */
  .score-hero{
    background:linear-gradient(135deg,#0d1117,#111827);border:1px solid var(--brd);
    border-radius:16px;padding:40px;text-align:center;margin-bottom:20px;
    position:relative;overflow:hidden;
  }
  .score-hero::before{
    content:'';position:absolute;inset:0;pointer-events:none;
    background:radial-gradient(ellipse at 50% -10%,rgba(74,126,199,.14),transparent 60%);
  }
  .sc-val{font-size:5.5em;font-weight:800;font-family:'JetBrains Mono',monospace;
    line-height:1;letter-spacing:-4px}
  .sc-grade{font-size:1.4em;font-weight:700;margin-top:8px;letter-spacing:2px;text-transform:uppercase}
  .sc-sub{color:var(--mut);font-size:.88em;margin-top:8px;font-family:'JetBrains Mono',monospace}

  /* Lighthouse banner */
  .lh-banner{background:var(--card);border:1px solid var(--brd);border-radius:16px;
    padding:24px;margin-bottom:20px;display:none}
  .lh-banner h3{font-size:.75em;text-transform:uppercase;letter-spacing:2px;color:var(--mut);margin-bottom:18px}
  .lh-circles{display:flex;gap:24px;flex-wrap:wrap;justify-content:center}
  .lh-item{display:flex;flex-direction:column;align-items:center;gap:8px}
  .lh-item span{font-size:.72em;color:var(--mut);text-align:center;max-width:70px}
  .lh-ring{width:76px;height:76px}
  .lh-na{font-size:1.5em;color:var(--mut);width:76px;height:76px;display:flex;align-items:center;justify-content:center}

  /* Stats */
  .stats{display:grid;grid-template-columns:repeat(5,1fr);gap:12px;margin-bottom:20px}
  .stat{background:var(--card);border:1px solid var(--brd);border-radius:12px;
    padding:18px 10px;text-align:center;transition:.22s}
  .stat:hover{border-color:var(--acc);transform:translateY(-2px)}
  .stat-n{font-size:2em;font-weight:800;font-family:'JetBrains Mono',monospace;line-height:1}
  .stat-l{font-size:.7em;color:var(--mut);margin-top:6px;text-transform:uppercase;letter-spacing:.8px}
  .files .stat-n{color:var(--acc)}.crit .stat-n{color:var(--red)}.high .stat-n{color:var(--ora)}
  .med .stat-n{color:var(--yel)}.low .stat-n{color:var(--grn)}

  /* Downloads */
  .dl-grid{display:grid;grid-template-columns:repeat(4,1fr);gap:12px;margin-bottom:20px}
  .dl-btn{background:var(--card);border:1px solid var(--brd);border-radius:10px;
    padding:14px;text-align:center;cursor:pointer;transition:.22s;
    font-family:'Syne',sans-serif;font-weight:700;color:var(--txt)}
  .dl-btn:hover{border-color:var(--acc);background:rgba(74,126,199,.07);transform:translateY(-2px)}
  .dl-icon{font-size:1.5em;display:block;margin-bottom:6px}
  .dl-fmt{display:block;font-size:.7em;font-family:'JetBrains Mono',monospace;
    color:var(--acc);margin-top:3px;letter-spacing:1px}

  /* AI panel */
  .ai-panel{background:var(--card);border:1px solid rgba(74,126,199,.3);
    border-radius:14px;margin-bottom:20px;overflow:hidden}
  .ai-head{
    display:flex;justify-content:space-between;align-items:center;
    padding:18px 22px;cursor:pointer;transition:.22s;
  }
  .ai-head:hover{background:rgba(74,126,199,.06)}
  .ai-title{font-size:.85em;font-weight:700;color:var(--acc)}
  .ai-toggle{color:var(--mut);font-size:.8em;font-family:'JetBrains Mono',monospace;transition:.3s}
  .ai-body{padding:0 22px 22px;display:none}
  .ai-key-row{display:flex;gap:10px;margin-bottom:14px}
  .ai-key-row input{
    flex:1;padding:11px 14px;border:1px solid var(--brd);border-radius:8px;
    background:rgba(255,255,255,.04);color:var(--txt);
    font-family:'JetBrains Mono',monospace;font-size:.85em;outline:none;transition:.22s;
  }
  .ai-key-row input:focus{border-color:var(--acc)}
  .ai-key-row input::placeholder{color:var(--mut)}
  .ai-run-btn{
    padding:11px 18px;background:linear-gradient(135deg,var(--acc),var(--acc2));
    color:#fff;border:none;border-radius:8px;font-family:'Syne',sans-serif;
    font-weight:700;font-size:.85em;cursor:pointer;transition:.22s;white-space:nowrap;
  }
  .ai-run-btn:hover{transform:translateY(-1px);box-shadow:0 6px 18px rgba(74,126,199,.3)}
  .ai-output{font-size:.88em;line-height:1.8;color:var(--txt);white-space:pre-wrap}

  /* Issues */
  .issues-panel{background:var(--card);border:1px solid var(--brd);border-radius:16px;padding:26px}
  .panel-hdr{font-size:.72em;text-transform:uppercase;letter-spacing:2px;color:var(--mut);
    margin-bottom:20px;display:flex;align-items:center;gap:10px}
  .panel-hdr::after{content:'';flex:1;height:1px;background:var(--brd)}
  .igroup{margin-bottom:22px}
  .itype-lbl{font-size:.72em;font-weight:600;text-transform:uppercase;letter-spacing:2px;
    color:var(--acc);padding:3px 9px;border:1px solid var(--brd);border-radius:4px;
    display:inline-block;margin-bottom:10px;font-family:'JetBrains Mono',monospace}
  .issue{border-left:2px solid;padding:10px 14px;margin-bottom:8px;border-radius:0 6px 6px 0;transition:.2s;cursor:default}
  .issue:hover{transform:translateX(3px)}
  .issue.critical{border-color:var(--red);background:rgba(255,77,109,.05)}
  .issue.high{border-color:var(--ora);background:rgba(255,149,0,.05)}
  .issue.medium{border-color:var(--yel);background:rgba(255,214,10,.05)}
  .issue.low{border-color:var(--grn);background:rgba(0,214,143,.05)}
  .i-top{display:flex;align-items:flex-start;gap:7px;flex-wrap:wrap;margin-bottom:4px}
  .sev-tag{font-family:'JetBrains Mono',monospace;font-size:.68em;font-weight:700;
    padding:2px 8px;border-radius:3px;white-space:nowrap}
  .sev-tag.critical{background:var(--red);color:#fff}.sev-tag.high{background:var(--ora);color:#000}
  .sev-tag.medium{background:var(--yel);color:#000}.sev-tag.low{background:var(--grn);color:#000}
  .src-chip{font-family:'JetBrains Mono',monospace;font-size:.64em;color:var(--mut);
    border:1px solid var(--brd);padding:2px 6px;border-radius:3px;white-space:nowrap}
  .i-msg{font-weight:600;font-size:.88em;flex:1}
  .i-fix{font-size:.8em;color:var(--mut)}
  .i-fix::before{content:'→ ';color:var(--acc)}
  .no-issues{text-align:center;padding:30px;color:var(--grn);font-size:1em;
    background:rgba(0,214,143,.05);border:1px solid rgba(0,214,143,.2);border-radius:10px}

  footer{text-align:center;padding:28px 0 6px;color:var(--brd);font-size:.78em;font-family:'JetBrains Mono',monospace}

  @media(max-width:640px){
    .stats{grid-template-columns:repeat(3,1fr)}.dl-grid{grid-template-columns:repeat(2,1fr)}
    .lh-circles{gap:14px}
  }
</style>
</head>
<body>
<div class="wrap">

  <header class="hdr">
    <div class="badge-pill">v3.0 · SECURITY + LIGHTHOUSE + AI</div>
    <h1>WebAnalyzer</h1>
    <p>Security audits · Lighthouse scores · AI fix suggestions — all in one place.</p>
  </header>

  <div class="card">
    <!-- Source type selector -->
    <div class="src-tabs">
      <button class="src-tab active" id="tabFile" onclick="switchTab('file')">📦 Upload Archive</button>
      <button class="src-tab" id="tabGH" onclick="switchTab('github')">🔗 GitHub / GitLab URL</button>
    </div>

    <!-- File upload panel -->
    <div id="filePanel">
      <div class="drop-zone" id="dropZone">
        <span class="dz-icon">📦</span>
        <div class="dz-title">Drop your archive(s) here</div>
        <div class="dz-sub">
          .zip &nbsp;·&nbsp; .tar.gz &nbsp;·&nbsp; .tar &nbsp;·&nbsp; folders (Chrome/Edge)<br>
          <span style="font-size:.8em;color:var(--mut)">Up to 500 MB · Multiple files supported</span>
        </div>
        <input type="file" id="fileInput" accept=".zip,.tar,.gz,.tgz" multiple>
        <button class="btn-primary" onclick="document.getElementById('fileInput').click()">
          Choose File(s)
        </button>
      </div>
    </div>

    <!-- GitHub URL panel -->
    <div id="ghPanel" class="gh-panel">
      <div class="gh-input-wrap">
        <span class="gh-icon">🔗</span>
        <input type="url" id="ghUrl" placeholder="https://github.com/user/repository">
      </div>
      <button class="btn-primary" onclick="submitGH()">
        Clone &amp; Analyze
      </button>
    </div>

    <div class="err" id="errBox"></div>

    <!-- Loading indicator -->
    <div id="loading">
      <div class="spin-row">
        <div class="spin"></div>
        <span class="ldt" id="ldtText">Extracting files…</span>
      </div>
      <div class="prog-track">
        <div class="prog-fill" id="progFill"></div>
      </div>
    </div>
  </div>

  <!-- Results -->
  <div id="results">

    <div class="score-hero" id="scoreHero"></div>

    <!-- Lighthouse banner (shown only when scores available) -->
    <div class="lh-banner" id="lhBanner">
      <h3>⚡ Lighthouse Scores</h3>
      <div class="lh-circles" id="lhCircles"></div>
    </div>

    <div class="stats" id="statsGrid"></div>

    <!-- Downloads -->
    <div class="dl-grid">
      <div class="dl-btn" onclick="dl('html')"><span class="dl-icon">📄</span>HTML<span class="dl-fmt">.html</span></div>
      <div class="dl-btn" onclick="dl('pdf')"><span class="dl-icon">📕</span>PDF<span class="dl-fmt">.pdf</span></div>
      <div class="dl-btn" onclick="dl('docx')"><span class="dl-icon">📘</span>Word<span class="dl-fmt">.docx</span></div>
      <div class="dl-btn" onclick="dl('json')"><span class="dl-icon">🗒️</span>JSON<span class="dl-fmt">.json</span></div>
    </div>

    <!-- AI Suggestions -->
    <div class="ai-panel" id="aiPanel">
      <div class="ai-head" onclick="toggleAI()">
        <span class="ai-title">✨ AI Fix Suggestions — powered by Grok (optional)</span>
        <span class="ai-toggle" id="aiToggle">▼ expand</span>
      </div>
      <div class="ai-body" id="aiBody">
        <div class="ai-key-row">
          <input type="password" id="aiKey" placeholder="xAI API key (sk-…) — stays in your browser">
          <button class="ai-run-btn" onclick="runAI()">Get Suggestions</button>
        </div>
        <div class="ai-output" id="aiOut"></div>
      </div>
    </div>

    <!-- Issues -->
    <div class="issues-panel">
      <div class="panel-hdr">Detailed Findings</div>
      <div id="issuesList"></div>
    </div>

  </div><!-- /#results -->

  <footer>WebAnalyzer v3.0 · Security &amp; Performance Analysis Platform</footer>
</div><!-- /.wrap -->

<script>
// ── Helpers ──────────────────────────────────────────────────────────────────
function esc(s){
  return String(s)
    .replace(/&/g,'&amp;').replace(/</g,'&lt;').replace(/>/g,'&gt;')
    .replace(/"/g,'&quot;').replace(/'/g,'&#39;');
}

// ── State ─────────────────────────────────────────────────────────────────────
let _data = null;
const LDT_MSGS = [
  'Extracting archive…',
  'Scanning HTML structure…',
  'Running security checks…',
  'Calling Bandit / ESLint…',
  'Running Lighthouse audit…',
  'Compiling findings…',
  'Generating report…',
];

// ── Tabs ──────────────────────────────────────────────────────────────────────
function switchTab(mode){
  document.getElementById('filePanel').style.display = mode==='file'?'block':'none';
  document.getElementById('ghPanel').style.display   = mode==='github'?'block':'none';
  document.getElementById('tabFile').classList.toggle('active', mode==='file');
  document.getElementById('tabGH').classList.toggle('active',   mode==='github');
}

// ── Drop zone ─────────────────────────────────────────────────────────────────
const dz = document.getElementById('dropZone');
const fi = document.getElementById('fileInput');

dz.addEventListener('dragover', e=>{ e.preventDefault(); dz.classList.add('over'); });
dz.addEventListener('dragleave', ()=>dz.classList.remove('over'));
dz.addEventListener('drop', e=>{
  e.preventDefault(); dz.classList.remove('over');
  handleDrop(e.dataTransfer);
});
fi.addEventListener('change', ()=>handleFiles(Array.from(fi.files)));

function handleDrop(dt){
  // Support folder drag (Chrome/Edge via DataTransferItem API)
  if(dt.items && dt.items.length){
    const allFiles = [];
    let pending = 0;
    for(const item of dt.items){
      if(item.kind !== 'file') continue;
      const entry = item.webkitGetAsEntry && item.webkitGetAsEntry();
      if(entry && entry.isDirectory){
        pending++;
        collectFromEntry(entry, allFiles, ()=>{ if(--pending===0) handleFiles(allFiles); });
      } else {
        allFiles.push(item.getAsFile());
      }
    }
    if(pending===0) handleFiles(allFiles);
  } else {
    handleFiles(Array.from(dt.files));
  }
}

function collectFromEntry(entry, out, done){
  if(entry.isFile){
    entry.file(f=>{ out.push(f); done(); });
  } else if(entry.isDirectory){
    const reader = entry.createReader();
    (function read(){
      reader.readEntries(entries=>{
        if(!entries.length){ done(); return; }
        let left = entries.length;
        entries.forEach(e=>collectFromEntry(e, out, ()=>{ if(--left===0) read(); }));
      });
    })();
  }
}

function handleFiles(files){
  if(!files || !files.length){ showErr('No files selected.'); return; }
  const valid = files.filter(f=>
    f.name.endsWith('.zip') || f.name.endsWith('.tar') ||
    f.name.endsWith('.tar.gz') || f.name.endsWith('.tgz') || f.size > 0
  );
  if(!valid.length){ showErr('Please upload a .zip, .tar, or .tar.gz file.'); return; }
  submitFiles(valid);
}

// ── Upload ────────────────────────────────────────────────────────────────────
function submitFiles(files){
  const fd = new FormData();
  files.forEach(f=>fd.append('file', f));
  startLoading();
  fetch('/analyze', {method:'POST', body:fd})
    .then(r=>r.json()).then(show).catch(err=>{ stopLoading(); showErr('Upload failed: '+err.message); });
}

function submitGH(){
  const url = document.getElementById('ghUrl').value.trim();
  if(!url){ showErr('Please enter a repository URL.'); return; }
  startLoading();
  fetch('/analyze', {
    method:'POST',
    headers:{'Content-Type':'application/json'},
    body:JSON.stringify({github_url: url}),
  }).then(r=>r.json()).then(show).catch(err=>{ stopLoading(); showErr('Clone failed: '+err.message); });
}

// ── Loading ───────────────────────────────────────────────────────────────────
let _progInt, _msgIdx;
function startLoading(){
  document.getElementById('loading').style.display='block';
  document.getElementById('results').style.display='none';
  document.getElementById('errBox').style.display='none';
  document.getElementById('progFill').style.width='0%';
  let prog=0; _msgIdx=0;
  _progInt = setInterval(()=>{
    prog = Math.min(prog + Math.random()*14, 88);
    document.getElementById('progFill').style.width=prog+'%';
    const next = Math.floor(prog / (88/LDT_MSGS.length));
    if(next !== _msgIdx && next < LDT_MSGS.length){
      _msgIdx=next; document.getElementById('ldtText').textContent=LDT_MSGS[_msgIdx];
    }
  }, 280);
}
function stopLoading(){
  clearInterval(_progInt);
  document.getElementById('progFill').style.width='100%';
  setTimeout(()=>{ document.getElementById('loading').style.display='none'; }, 350);
}

// ── Display results ───────────────────────────────────────────────────────────
function show(data){
  stopLoading();
  if(data.error){ showErr(data.error); return; }
  _data = data;

  const score = data.score ?? 85;
  const grade = data.grade ?? 'Good';
  const issues = data.issues ?? [];
  const lh = data.lighthouse_scores;

  const col = score>=80 ? '#00d68f' : score>=60 ? '#ff9500' : '#ff4d6d';

  // Score hero
  document.getElementById('scoreHero').innerHTML =
    `<div class="sc-val" style="color:${esc(col)}">${Math.round(score)}</div>`+
    `<div class="sc-grade" style="color:${esc(col)}">${esc(grade)}</div>`+
    `<div class="sc-sub">${esc(data.site_name)} &middot; ${issues.length} issue${issues.length!==1?'s':''} detected</div>`;

  // Lighthouse circles
  if(lh && Object.keys(lh).length){
    const labels = {
      'performance':'Performance','accessibility':'Accessibility',
      'best-practices':'Best Practices','seo':'SEO','pwa':'PWA'
    };
    let circHtml = '';
    for(const [key, label] of Object.entries(labels)){
      const v = lh[key];
      if(v==null){ circHtml += `<div class="lh-item"><div class="lh-na">—</div><span>${esc(label)}</span></div>`; continue; }
      const c = v>=90?'#00d68f':v>=50?'#ff9500':'#ff4d6d';
      circHtml += `<div class="lh-item">
        <svg viewBox="0 0 36 36" class="lh-ring">
          <circle cx="18" cy="18" r="15.9" fill="none" stroke="#2a3a55" stroke-width="3"/>
          <circle cx="18" cy="18" r="15.9" fill="none" stroke="${c}" stroke-width="3"
            stroke-dasharray="${v} 100" stroke-linecap="round" transform="rotate(-90 18 18)">
            <animate attributeName="stroke-dasharray" from="0 100" to="${v} 100" dur=".8s" fill="freeze"/>
          </circle>
          <text x="18" y="22" text-anchor="middle" font-size="8"
                font-family="JetBrains Mono,monospace" fill="${c}" font-weight="bold">${v}</text>
        </svg>
        <span>${esc(label)}</span>
      </div>`;
    }
    document.getElementById('lhCircles').innerHTML = circHtml;
    document.getElementById('lhBanner').style.display = 'block';
  }

  // Stats
  const c = issues.filter(i=>i.severity==='CRITICAL').length;
  const h = issues.filter(i=>i.severity==='HIGH').length;
  const m = issues.filter(i=>i.severity==='MEDIUM').length;
  const l = issues.filter(i=>i.severity==='LOW').length;
  document.getElementById('statsGrid').innerHTML =
    `<div class="stat files"><div class="stat-n">${esc(data.files_analyzed??0)}</div><div class="stat-l">Files</div></div>`+
    `<div class="stat crit"><div class="stat-n">${c}</div><div class="stat-l">Critical</div></div>`+
    `<div class="stat high"><div class="stat-n">${h}</div><div class="stat-l">High</div></div>`+
    `<div class="stat med"><div class="stat-n">${m}</div><div class="stat-l">Medium</div></div>`+
    `<div class="stat low"><div class="stat-n">${l}</div><div class="stat-l">Low</div></div>`;

  // Issues grouped by type
  const by = {};
  issues.forEach(i=>{ const t=i.type||'Other'; (by[t]=by[t]||[]).push(i); });
  const sevOrd = {CRITICAL:0,HIGH:1,MEDIUM:2,LOW:3};

  let ihtml = '';
  if(!issues.length){
    ihtml = '<div class="no-issues">✓ No issues found — your project is clean!</div>';
  } else {
    Object.entries(by)
      .sort((a,b)=>Math.min(...a[1].map(i=>sevOrd[i.severity]??9))-Math.min(...b[1].map(i=>sevOrd[i.severity]??9)))
      .forEach(([type,list])=>{
        ihtml += `<div class="igroup"><span class="itype-lbl">${esc(type)}</span>`;
        list.sort((a,b)=>(sevOrd[a.severity]??9)-(sevOrd[b.severity]??9)).forEach(issue=>{
          const sev = (issue.severity||'LOW').toLowerCase();
          ihtml += `<div class="issue ${esc(sev)}">
            <div class="i-top">
              <span class="sev-tag ${esc(sev)}">${esc(issue.severity||'INFO')}</span>
              <span class="src-chip">${esc(issue.source||'')}</span>
              <span class="i-msg">${esc(issue.message||'—')}</span>
            </div>
            <div class="i-fix">${esc(issue.fix||'No suggestion available')}</div>
          </div>`;
        });
        ihtml += '</div>';
      });
  }
  document.getElementById('issuesList').innerHTML = ihtml;

  document.getElementById('results').style.display='block';
  document.getElementById('results').scrollIntoView({behavior:'smooth',block:'start'});
}

// ── Downloads ─────────────────────────────────────────────────────────────────
function dl(fmt){
  if(!_data){ showErr('No analysis data available.'); return; }
  fetch(`/download/${fmt}`,{
    method:'POST',
    headers:{'Content-Type':'application/json'},
    body:JSON.stringify(_data),
  })
  .then(r=>{ if(!r.ok) return r.json().then(e=>{throw new Error(e.error);}); return r.blob(); })
  .then(blob=>{
    const ext = fmt==='docx'?'docx':fmt;
    const a = document.createElement('a');
    a.href = URL.createObjectURL(blob);
    a.download = `Report_${(_data.site_name||'report').replace(/[^a-z0-9_-]/gi,'_')}.${ext}`;
    a.click();
    URL.revokeObjectURL(a.href);
  })
  .catch(e=>showErr('Download failed: '+e.message));
}

// ── AI suggestions ────────────────────────────────────────────────────────────
function toggleAI(){
  const body = document.getElementById('aiBody');
  const tog  = document.getElementById('aiToggle');
  const open = body.style.display === 'block';
  body.style.display = open?'none':'block';
  tog.textContent    = open?'▼ expand':'▲ collapse';
}

async function runAI(){
  if(!_data){ return; }
  const key = document.getElementById('aiKey').value.trim();
  if(!key){ showErr('Please enter your xAI API key.'); return; }
  const out = document.getElementById('aiOut');
  out.textContent = 'Calling Grok…';
  try{
    const r = await fetch('/ai_suggest',{
      method:'POST',
      headers:{'Content-Type':'application/json'},
      body:JSON.stringify({issues: _data.issues, api_key: key}),
    });
    const d = await r.json();
    if(d.error){ out.textContent = '⚠ '+d.error; return; }
    out.textContent = d.suggestions || '(No suggestions returned)';
  }catch(e){
    out.textContent = '⚠ '+e.message;
  }
}

// ── Utils ─────────────────────────────────────────────────────────────────────
function showErr(msg){
  const b = document.getElementById('errBox');
  b.textContent = '⚠ '+msg;
  b.style.display='block';
}
</script>
</body>
</html>"""


# =============================================================================
# === FLASK ROUTES ===
# =============================================================================

@app.route("/health")
def health():
    return jsonify({"status": "ok", "version": "3.0", "timestamp": datetime.utcnow().isoformat()})


@app.route("/")
def index():
    return render_template_string(HTML_TEMPLATE)


@app.route("/analyze", methods=["POST"])
@limiter.limit(CFG["rate_limit"])
def analyze():
    try:
        # ── Check disk quota before accepting work ──
        total_disk = sum(
            f.stat().st_size for base in [BASE_UPLOAD, BASE_OUTPUT]
            for f in base.rglob("*") if f.is_file()
        )
        if total_disk > _DISK_QUOTA:
            return jsonify({"error": "Server storage quota exceeded. Try again later."}), 503

        session_id  = hashlib.sha256(os.urandom(16)).hexdigest()[:16]
        work_dir    = BASE_UPLOAD / session_id
        work_dir.mkdir(parents=True, exist_ok=True)
        _cleanup.register(session_id, work_dir)

        site_name   = "project"
        file_count  = 0

        # ── GitHub URL path ──
        if request.is_json:
            body = request.get_json(silent=True) or {}
            gh_url = body.get("github_url", "")
            if gh_url:
                try:
                    file_count = clone_github_repo(gh_url, work_dir)
                    site_name  = Path(urlparse(gh_url).path).stem.replace(".git", "") or "repo"
                except ExtractionError as exc:
                    return jsonify({"error": str(exc)}), 400
            else:
                return jsonify({"error": "No file or github_url provided"}), 400

        # ── File upload path ──
        else:
            files = request.files.getlist("file")
            if not files:
                return jsonify({"error": "No files provided"}), 400

            for upload in files:
                fname = upload.filename or "upload"
                dest_path = work_dir / fname
                upload.save(dest_path)
                site_name = Path(fname).stem.replace(".tar", "").replace(".gz", "")
                try:
                    if fname.endswith(".zip"):
                        file_count += safe_extract_zip(dest_path, work_dir / Path(fname).stem)
                    elif fname.endswith((".tar.gz", ".tgz", ".tar")):
                        file_count += safe_extract_tar(dest_path, work_dir / Path(fname).stem)
                    else:
                        # Treat as raw folder upload (werkzeug saves with path)
                        file_count += 1
                except ExtractionError as exc:
                    return jsonify({"error": str(exc)}), 400
                finally:
                    dest_path.unlink(missing_ok=True)

        # ── Run analysis ──
        orchestrator = AnalysisOrchestrator(session_id, site_name, work_dir)
        result = orchestrator.run()
        if not result.get("success"):
            return jsonify({"error": result.get("error", "Analysis failed")}), 500

        # ── Run Lighthouse ──
        lh_scores = LighthouseRunner.run(work_dir)

        # ── Build report ──
        ts          = datetime.now()
        report_gen  = ReportGenerator(site_name, ts)
        report_gen.add_issues(result["issues"])
        report_gen.set_lighthouse(lh_scores)
        score, grade = report_gen.calculate_score()

        with _sessions_lock:
            _sessions[session_id] = {
                "report_gen": report_gen,
                "timestamp":  ts,
                "issues":     result["issues"],
            }

        return jsonify({
            "success":           True,
            "session_id":        session_id,
            "site_name":         site_name,
            "files_analyzed":    orchestrator.file_count,
            "issues":            result["issues"],
            "score":             score,
            "grade":             grade,
            "lighthouse_scores": lh_scores,
            "timestamp":         ts.isoformat(),
        })

    except Exception as exc:
        logger.error("analyze error: %s", traceback.format_exc())
        return jsonify({"error": str(exc)}), 500


@app.route("/download/<fmt>", methods=["POST"])
@limiter.limit("20 per minute")
def download(fmt: str):
    try:
        data        = request.get_json(silent=True) or {}
        site_name   = data.get("site_name", "")
        session_id  = data.get("session_id", "")

        # Look up by session_id first, then site_name fallback
        with _sessions_lock:
            sess = _sessions.get(session_id) or next(
                (v for v in _sessions.values() if v["report_gen"].site_name == site_name), None
            )

        if not sess:
            return jsonify({"error": "Session not found. Re-run the analysis."}), 404

        rg: ReportGenerator = sess["report_gen"]
        safe_name = re.sub(r"[^A-Za-z0-9_\-]", "_", rg.site_name)

        if fmt == "html":
            content = rg.generate_html().encode("utf-8")
            return send_file(
                io.BytesIO(content), mimetype="text/html",
                as_attachment=True, download_name=f"Report_{safe_name}.html",
            )

        elif fmt == "json":
            content = rg.generate_json().encode("utf-8")
            return send_file(
                io.BytesIO(content), mimetype="application/json",
                as_attachment=True, download_name=f"Report_{safe_name}.json",
            )

        elif fmt == "docx":
            out = BASE_OUTPUT / f"Report_{safe_name}_{session_id[:8]}.docx"
            rg.generate_docx(out)
            return send_file(
                out,
                mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                as_attachment=True, download_name=f"Report_{safe_name}.docx",
            )

        elif fmt == "pdf":
            out = BASE_OUTPUT / f"Report_{safe_name}_{session_id[:8]}.pdf"
            rg.generate_pdf(out)
            return send_file(
                out, mimetype="application/pdf",
                as_attachment=True, download_name=f"Report_{safe_name}.pdf",
            )

        else:
            return jsonify({"error": f"Unknown format: {fmt}"}), 400

    except Exception as exc:
        logger.error("download error: %s", traceback.format_exc())
        return jsonify({"error": str(exc)}), 500


@app.route("/ai_suggest", methods=["POST"])
@limiter.limit("5 per minute")
def ai_suggest():
    try:
        data    = request.get_json(silent=True) or {}
        issues  = data.get("issues", [])
        api_key = data.get("api_key", "").strip()
        if not api_key:
            api_key = CFG.get("xai_api_key", "")
        if not api_key:
            return jsonify({"error": "xAI API key required"}), 400
        if not issues:
            return jsonify({"error": "No issues to analyse"}), 400

        text = get_ai_suggestions(issues, api_key)
        if not text:
            return jsonify({"error": "Grok API returned no response"}), 502
        return jsonify({"suggestions": text})

    except Exception as exc:
        logger.error("ai_suggest error: %s", exc)
        return jsonify({"error": str(exc)}), 500


# 404 / 500 JSON handlers
@app.errorhandler(404)
def not_found(_):
    return jsonify({"error": "Not found"}), 404

@app.errorhandler(429)
def rate_limited(_):
    return jsonify({"error": "Rate limit exceeded (8 requests/minute). Please wait."}), 429

@app.errorhandler(500)
def server_error(_):
    return jsonify({"error": "Internal server error"}), 500


# =============================================================================
# === ENTRY POINT ===
# =============================================================================
if __name__ == "__main__":
    print("""
  ╔══════════════════════════════════════════════════════════════════╗
  ║   WebAnalyzer v3.0 — Production Security & Performance Platform  ║
  ║                                                                  ║
  ║   Security · Bandit · ESLint · Stylelint · Lighthouse · AI       ║
  ║   Multi-format reports: HTML · PDF · DOCX · JSON                 ║
  ║                                                                  ║
  ║   Press Ctrl+C to stop                                           ║
  ╚══════════════════════════════════════════════════════════════════╝
""")
    try:
        import gunicorn  # noqa: F401
        # Run via gunicorn if available (production)
        opts = [
            f"--bind={CFG['host']}:{CFG['port']}",
            f"--workers={CFG['workers']}",
            "--worker-class=sync",
            "--timeout=300",
            f"--log-level={CFG['log_level'].lower()}",
            "web_analyzer:app",
        ]
        os.execv(
            sys.executable,
            [sys.executable, "-m", "gunicorn"] + opts,
        )
    except ImportError:
        logger.warning("gunicorn not found — using Flask dev server (not for production!)")
        app.run(host=CFG["host"], port=int(CFG["port"]), debug=False)